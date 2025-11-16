"""Utility functions for parsing resume files (PDF, DOCX)."""
import io
from typing import Optional
import PyPDF2
from docx import Document

async def parse_pdf(file_content: bytes) -> Optional[str]:
    """Extract text from PDF file."""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []
        
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return '\n'.join(text_parts) if text_parts else None
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


async def parse_docx(file_content: bytes) -> Optional[str]:
    """Extract text from DOCX file."""
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts) if text_parts else None
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


async def parse_resume_file(file_content: bytes, file_type: str) -> Optional[str]:
    """Parse resume file based on file type."""
    if file_type == 'application/pdf' or file_type == 'pdf':
        return await parse_pdf(file_content)
    elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword', 'docx', 'doc']:
        return await parse_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

